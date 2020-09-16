from pdfCollectTest0728.configCBIM.fileConfig import fileConfig
import docx

class wordData:

    def __init__(self):
        self.file = fileConfig().fileAddress["word"]
        self.wordDatas = self.openWord()
        self.wordList = self.cutWord()

    def openWord(self):
        # 读取word文本
        wordDatas = docx.Document(self.file)
        return wordDatas

    def cutWord(self):
        # 将word文本分割为每段
        wordList=[]
        print("段落数:"+str(len(self.wordDatas.paragraphs)))#段落数为13，每个回车隔离一段
        # 输出每一段的内容
        for para in self.wordDatas.paragraphs:
            # print(para.text)
            if len(para.text) > 0:
                wordList.append(para.text)
            else:
                wordList.append("*无法识别*")
        # 输出段落编号及段落内容
        # for i in range(len(wordList)):
        #     print("第"+str(i)+"段的内容是："+wordList[i])
        return wordList



# wordData = wordData()





